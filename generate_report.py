from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

def create_report():
    doc = Document()
    
    # --- PAGE 1: COVER ---
    title = doc.add_heading('BÁO CÁO CUỐI NGÀY', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Thực tập sinh Frontend - Thực hành xây dựng giao diện trang Giới thiệu và trang Đối tác', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n')
    info_p = doc.add_paragraph()
    info_p.add_run('Họ và tên: ').bold = True
    info_p.add_run('[Điền tên của bạn]\n')
    info_p.add_run('Vị trí: ').bold = True
    info_p.add_run('Thực tập sinh Frontend\n')
    info_p.add_run('Ngày thực tập: ').bold = True
    info_p.add_run('20/07/2026\n')
    info_p.add_run('Dự án: ').bold = True
    info_p.add_run('Website VUSTIC JSC\n')
    
    doc.add_page_break()
    
    # --- PAGE 2: MỤC TIÊU & TỔNG QUAN ---
    doc.add_heading('PHẦN 1: TỔNG QUAN VÀ MỤC TIÊU CÔNG VIỆC', level=1)
    
    doc.add_heading('1.1. Mục tiêu công việc trong ngày', level=2)
    doc.add_paragraph('Trong ngày làm việc 20/07/2026, mục tiêu chính được đề ra là tiếp tục hoàn thiện giao diện người dùng (UI) cho website VUSTIC. Trọng tâm của ngày hôm nay được đặt vào hai trang vệ tinh quan trọng là "Giới thiệu" (About Page) và "Đối tác" (Partners Page). Việc xây dựng hai trang này không chỉ yêu cầu tính thẩm mỹ, bám sát bản thiết kế gốc mà còn đòi hỏi sự tối ưu về mặt kiến trúc mã nguồn (code architecture) và hiệu năng.')
    doc.add_paragraph('Các mục tiêu cụ thể bao gồm:')
    doc.add_paragraph('- Khắc phục các cảnh báo lỗi (Deprecation Warning) của hệ thống biên dịch SCSS liên quan đến @import.', style='List Bullet')
    doc.add_paragraph('- Hoàn thiện cấu trúc trang Giới thiệu, bao gồm cả hiệu ứng hình ảnh động và căn chỉnh bố cục đáp ứng (responsive).', style='List Bullet')
    doc.add_paragraph('- Khởi tạo và xây dựng từ đầu trang Đối tác, thiết lập hệ thống quản lý dữ liệu đối tác tĩnh.', style='List Bullet')
    doc.add_paragraph('- Áp dụng các kỹ thuật nâng cao của ReactJS như Hooks (useState, useEffect, useRef) để xử lý logic hiển thị.', style='List Bullet')
    
    doc.add_heading('1.2. Tổng quan tiến độ', level=2)
    doc.add_paragraph('Về cơ bản, tiến độ công việc trong ngày đã hoàn thành xuất sắc 100% mục tiêu đề ra. Cả hai trang Giới thiệu và Đối tác đều đã được tích hợp thành công vào Router chính của dự án, có thể truy cập mượt mà. Hệ thống UI đã được kiểm tra trên cả giao diện Desktop và Mobile, đảm bảo tính responsive.')
    
    doc.add_paragraph('\n' * 5)
    doc.add_page_break()
    
    # --- PAGE 3: TRANG GIỚI THIỆU ---
    doc.add_heading('PHẦN 2: CHI TIẾT CÔNG VIỆC - TRANG GIỚI THIỆU', level=1)
    
    doc.add_heading('2.1. Cập nhật SCSS và Khắc phục Deprecation Warning', level=2)
    doc.add_paragraph('Bắt đầu ngày làm việc, hệ thống ghi nhận lỗi cảnh báo từ Sass liên quan đến việc sử dụng `@import`. Cụ thể, Sass đang dần loại bỏ `@import` và thay thế bằng `@use` hoặc `@forward` để quản lý module tốt hơn. Tôi đã tiến hành rà soát file global.scss và các file component liên quan, chuyển đổi cú pháp sang `@use` và đặt lên vị trí đầu tiên của tệp tin. Việc này giúp dự án tương thích tốt hơn với phiên bản Dart Sass 3.0.0 sắp tới và loại bỏ hoàn toàn các cảnh báo phiền nhiễu khi chạy lệnh npm run dev.')
    
    doc.add_heading('2.2. Xây dựng AboutHero và Fix Layout', level=2)
    doc.add_paragraph('Component AboutHero được tái cấu trúc để đồng bộ phong cách với trang chủ (HomePage). Tôi đã:')
    doc.add_paragraph('- Tái sử dụng kỹ thuật CSS Animation (scrollMarquee) để tạo dải logo đối tác chạy trượt ngang dưới chân banner.', style='List Bullet')
    doc.add_paragraph('- Xử lý lỗi layout khi header che khuất nội dung banner bằng cách bổ sung thuộc tính margin-top: 108px cho thẻ bọc ngoài, giúp đẩy nội dung xuống dưới vùng an toàn của Navigation bar trong suốt.', style='List Bullet')
    
    doc.add_heading('2.3. Hiệu ứng Intersection Observer cho Vision & Mission', level=2)
    doc.add_paragraph('Để tăng tính tương tác, tôi đã tích hợp Web API IntersectionObserver vào trang Tầm nhìn & Sứ mệnh. Bằng cách viết một component bao bọc (wrapper) mang tên AnimatedRow kết hợp với useRef và useEffect:')
    doc.add_paragraph('- Hệ thống sẽ liên tục theo dõi vị trí cuộn chuột của người dùng.', style='List Bullet')
    doc.add_paragraph('- Khi một khối nội dung lọt vào vùng nhìn thấy (viewport) khoảng 20%, một class CSS chứa hiệu ứng slide-in-left sẽ được gắn vào.', style='List Bullet')
    doc.add_paragraph('- Kết quả là các khối màu xanh, xanh đậm và đỏ lần lượt trượt từ trái sang và hiện rõ dần lên, tạo cảm giác vô cùng mượt mà và chuyên nghiệp.', style='List Bullet')
    
    doc.add_heading('2.4. Sửa lỗi hiển thị hình ảnh Nhân sự (Organization Chart)', level=2)
    doc.add_paragraph('Trong phần Sơ đồ tổ chức, các hình ảnh chân dung lãnh đạo gặp tình trạng bị cắt mất phần đỉnh đầu do thuộc tính object-fit: cover canh giữa tự động. Vấn đề này đã được xử lý triệt để bằng cách bổ sung thuộc tính object-position: top center, giúp trình duyệt ưu tiên giữ lại phần phía trên của bức ảnh, đảm bảo chân dung luôn hiển thị trọn vẹn khuôn mặt.')
    
    doc.add_page_break()
    
    # --- PAGE 4: TRANG ĐỐI TÁC ---
    doc.add_heading('PHẦN 3: CHI TIẾT CÔNG VIỆC - TRANG ĐỐI TÁC', level=1)
    
    doc.add_heading('3.1. Thiết kế UI/UX Component PartnersHero', level=2)
    doc.add_paragraph('Trang Đối tác được thiết kế với một Hero Banner nổi bật. Component PartnersHero được xây dựng kế thừa từ cấu trúc chung, nhưng sử dụng nền màu đỏ và thông điệp riêng: "VUSTIC mang sứ mệnh kết nối thương mại, đầu tư...". Dải băng chuyền đối tác (marquee) cũng được tích hợp ngay lập tức để duy trì nhận diện thương hiệu nhất quán trên toàn website.')
    
    doc.add_heading('3.2. Tổ chức Service lưu trữ Dữ liệu', level=2)
    doc.add_paragraph('Thay vì viết dữ liệu cứng (hard-code) trực tiếp vào file giao diện, tôi đã chủ động tạo ra một tệp tin dịch vụ tại src/services/PartnersList.js. Tệp tin này xuất ra ba hằng số quan trọng:')
    doc.add_paragraph('1. partnerCategories: Mảng chứa tên các danh mục (Bán lẻ, Tài chính, Sản xuất...).', style='List Bullet')
    doc.add_paragraph('2. featuredPartner: Object chứa thông tin đầy đủ của đối tác chiến lược nổi bật (Tập đoàn Cafe An Thái).', style='List Bullet')
    doc.add_paragraph('3. partners: Mảng chứa danh sách các đối tác còn lại, mỗi đối tác đều có thuộc tính category tương ứng.', style='List Bullet')
    doc.add_paragraph('Việc tách biệt này giúp codebase sạch sẽ hơn (Clean Code) và tạo tiền đề thuận lợi cho việc kết nối với API Backend (như Fetch hoặc Axios) trong tương lai.')
    
    doc.add_heading('3.3. Xử lý Logic Hiển thị & Lọc Dữ liệu (Filtering)', level=2)
    doc.add_paragraph('Tính năng phức tạp nhất trong ngày là bộ lọc đối tác. Tôi đã sử dụng hook useState để lưu trữ trạng thái activeTab. Khi người dùng nhấn vào các nút danh mục, state này được cập nhật.')
    doc.add_paragraph('Kỹ thuật được áp dụng bao gồm:')
    doc.add_paragraph('- Conditional Rendering (Toán tử &&): Chỉ render khối đối tác nổi bật (featuredPartner) khi người dùng đang ở tab "Tất cả".', style='List Bullet')
    doc.add_paragraph('- Lọc mảng (Array.filter): Sử dụng logic p.category === activeTab || activeTab === "Tất cả" để tự động rút gọn danh sách đối tác theo đúng nhu cầu.', style='List Bullet')
    doc.add_paragraph('- Lặp phần tử (Array.map): Tự động sinh ra các thẻ HTML partner-card dựa trên danh sách đã lọc, giúp tiết kiệm hàng trăm dòng code lặp lại.', style='List Bullet')
    
    doc.add_page_break()
    
    # --- PAGE 5: KẾT LUẬN ---
    doc.add_heading('PHẦN 4: KIẾN THỨC VÀ KỸ NĂNG ĐƯỢC CỦNG CỐ', level=1)
    
    doc.add_paragraph('Trong quá trình thực hiện các task trên, tôi đã có cơ hội vận dụng và củng cố sâu sắc các kiến thức sau:')
    doc.add_paragraph('1. ReactJS Cơ bản & Nâng cao: Sự khác biệt giữa biến thông thường và state, lý do không được truy xuất trực tiếp các thuộc tính không tồn tại trong map (ví dụ lỗi p.activeTab). Cách sử dụng className động với Template Literals (dấu backtick).', style='List Bullet')
    doc.add_paragraph('2. CSS/SCSS Styling: Làm chủ Flexbox và CSS Grid Layout để xếp chồng các thành phần, chia cột responsive. Kỹ thuật tạo animation với @keyframes.', style='List Bullet')
    doc.add_paragraph('3. Tư duy Component-based: Chia nhỏ UI thành các khối độc lập (AnimatedRow, SectionTitle, ProjectCard) dù ở các file khác nhau, nhưng vẫn duy trì tính đồng nhất.', style='List Bullet')
    
    doc.add_heading('PHẦN 5: KHÓ KHĂN VÀ KẾ HOẠCH TIẾP THEO', level=1)
    
    doc.add_heading('5.1. Khó khăn gặp phải', level=2)
    doc.add_paragraph('- Việc kiểm soát không gian khi thẻ Navigation là trong suốt (transparent) đòi hỏi phải tính toán kỹ lưỡng padding/margin cho các thẻ Hero Banner để không bị lệch bố cục.')
    doc.add_paragraph('- Việc gắn API IntersectionObserver vào React yêu cầu hiểu rõ về vòng đời component (đặc biệt là việc gọi disconnect() trong phần cleanup function của useEffect) để tránh rò rỉ bộ nhớ (memory leak).')
    
    doc.add_heading('5.2. Kế hoạch cho phiên làm việc tới', level=2)
    doc.add_paragraph('- Tiến hành cập nhật hình ảnh thật thay thế cho các ảnh placeholder hiện tại trong thư mục assets.', style='List Bullet')
    doc.add_paragraph('- Refactor lại trang HomePage để tái sử dụng lại các component chung như SectionTitle, ProjectCard nhằm dọn dẹp mã nguồn rác.', style='List Bullet')
    doc.add_paragraph('- Tích hợp tính năng gửi Form liên hệ trong trang ContactPage.', style='List Bullet')
    doc.add_paragraph('- Rà soát lỗi responsive tổng thể trên các kích thước màn hình nhỏ (Tablet, Mobile) trước khi tiến hành tối ưu hóa SEO (Meta tags).', style='List Bullet')
    
    doc.save("Bao_Cao_Thuc_Tap_Frontend.docx")

if __name__ == "__main__":
    create_report()
